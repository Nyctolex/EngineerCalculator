import math
import pandas as pd
import lab_calculator
import pyperclip
import pandas as pd
import sympy as sp
from sympy.printing.mathml import mathml
import re
import eddingtone_bridge
from docx import Document
from docx.shared import Inches
import os

_OUT_FOLDER = r"C:\Users\chris-chen\Documents\TAU\Year B\Semester A\lab\דוחות\צמיגות\output"
_KEYWORDS_FILTER = []

CLIPBOARD_COPY = 0
#Should the returned data should be copied to clipboard or not.
PRINT_VALUES = False
#Should the returned data should be printed or not.
    

def print_or_copy(data, print_values=PRINT_VALUES, copy_values=CLIPBOARD_COPY):
    """
    Choosing the correct action according to the global values.
    """
    if PRINT_VALUES:
        #print(data)
        pass
    if CLIPBOARD_COPY:
        pyperclip.copy(data)

def relative_error_equaion(sign, res, er):
    rel_er = lab_calculator.relative_error(res, er)
    res_for = "Δ{sign} /{sign} = {rel_er}%".format(sign=sign, rel_er = rel_er)
    print_or_copy(res_for)
    return res_for

def get_val_and_error_format(sign, expression, error_dict, values_dict):
    val, er = lab_calculator.get_val_and_error(expression, error_dict, values_dict)
    res_format = "{sign} = {val} +- {er}".format(sign=sign, val=val, er=er)
    print_or_copy(res_format)
    return(res_format)

def format_eq(equation):
    while ("  " in equation):
        equation = equation.replace("  ", " ")
    while ("--" in equation) or ("- -" in equation):
        equation = equation.replace("--", "+")
        equation = equation.replace("- -", "+")
    while ("++" in equation) or ("+ +" in equation):
        equation = equation.replace("++", "+")
        equation = equation.replace("+ +", "+")
    return equation

def total_error(errors_lst):
    total = 0.0
    for er in errors_lst:
        total += er**2
    total = math.sqrt(total)
    return lab_round(er)

def n_sigma_format(theory, theory_error, value, error_value):
    res = lab_calculator.calculate_n_sigma(theory, theory_error, value, error_value)
    res = lab_calculator.lab_round(res)
    txt = "N_σ=|({theo}−{val})|/√(({theo_error})^2+({val_er})^2 )={res}".format(theo=theory, val=value, theo_error=theory_error, val_er=error_value, res=res)
    txt = format_eq(txt)
    print_or_copy(txt)
    return txt

def replace_brackets(st):
    st = st.replace('{', '(')
    st = st.replace('}', ')')
    return st

def render_factorial(expression):
    expression = expression.replace('\frac', '\\frac')
    print(expression)
    while (len(re.findall(r'\\frac{.+?}{.+?}', expression)) != 0):
        factors = re.search(r'\\frac{.+?}{.+?}', expression)
        sub_ex = factors.group(0)
        part1 = re.search(r'\\frac{.+?}{', sub_ex).group(0)[len('\\frac'):-1]
        part2 = sub_ex[len('\\frac{')+len(part1)-1:]
        part1 = replace_brackets(part1)
        part2 = replace_brackets(part2)
        new_ex = "{part1}/{part2}".format(part1=part1, part2=part2)
        expression = expression.replace(sub_ex, new_ex)
    return expression

def spympy_to_word(expretion):
    res = sp.latex(expretion)
    res = render_factorial(res)
    res = replace_brackets(res)
    print_or_copy(res)
    return res

class ducument_analyser():
    def __init__(self, headFolder, formulas_dict, value_dict, error_dict, keyword_filter=[]):
        self.headFolder = headFolder
        self.keyword_filter = keyword_filter
        self.data_obj = eddingtone_bridge.folder_scanner(self.headFolder, self.keyword_filter)
        self.folrmulas_dict = formulas_dict
        self.value_dict = value_dict
        self.error_dict = error_dict
        self.document = Document()
    
    def insert_table(self, df):
        t = self.document.add_table(df.shape[0]+1, df.shape[1])
        # add the header rows.
        for j in range(df.shape[-1]):
            t.cell(0,j).text = df.columns[j]
        # add the rest of the data frame
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                t.cell(i+1,j).text = str(df.values[i,j])

    @staticmethod
    def insert_new_line(run, data):
        run.add_text(data)
        run.add_break()
        run.add_break()
    
    def get_graphs_name(self, sub_folder):
        fit = residuals = None
        files = os.listdir(self.headFolder + "\\" + sub_folder)
        for img_name in files:
            if 'fitting' in img_name:
                fit = self.headFolder + "\\" + sub_folder + "\\" + img_name
            if 'residuals' in img_name:
                residuals = self.headFolder + "\\" + sub_folder + "\\" + img_name
        return fit, residuals
    
    def generate_document(self, relative_dict={}):
        filtered_folder = self.data_obj.load_folders().copy()
        for sub_folder in filtered_folder:
            self.data_obj.keyword_filter = [sub_folder]
            self.data_obj.load_folders()
            self.data_obj.load_plot_files()
            rel_er_list = []
            for key in self.folrmulas_dict:
                if key in sub_folder:
                    expression, product_symbol, extract_value, rel_er_list = self.folrmulas_dict[sub_folder]
                    no_value = int(re.search(r'[0-9]+', spympy_to_word(extract_value)).group())
                    if rel_er_list == []:
                        rel_er_list = [no_value]
            a_no = self.data_obj.get_max_a_in_plot(self.data_obj.plot_files[0])
            p = self.document.add_paragraph()
            r = p.add_run()
            r.add_text('Analyse folder: ' + sub_folder)
            self.insert_table(self.data_obj.plot_to_table(a_no, rel_er_lst=rel_er_list))
            p = self.document.add_paragraph()
            fit, residuals = self.get_graphs_name(sub_folder)
            r.add_picture(fit)
            r.add_picture(residuals)
            expression = None
            if expression is not None:
                sign = spympy_to_word(product_symbol)
                error_sign = spympy_to_word(self.error_dict[product_symbol])
                r = p.add_run()
                r.add_break()
                r.add_break()
                self.insert_new_line(r, "Expresstion:")
                self.insert_new_line(r, sign + " = " + spympy_to_word(expression))
                self.insert_new_line(r, "Expresstion error:")
                self.insert_new_line(r, error_sign + " = " + spympy_to_word(lab_calculator.get_error_formula(expression, self.error_dict)))
               
                self.value_dict[extract_value] = self.data_obj.load_a(no_value)[0]
                self.value_dict[self.error_dict[extract_value]] = self.data_obj.load_d(no_value)[0]
                res, er = lab_calculator.get_val_and_error(expression, self.error_dict, self.value_dict)
                output = get_val_and_error_format(sign, expression, self.error_dict, self.value_dict)
                self.insert_new_line(r, "Product Value:")
                self.insert_new_line(r, output)
                self.insert_new_line(r, "Relative Error:")
                output = relative_error_equaion(sign, res, er)
                self.insert_new_line(r, output)
                self.insert_new_line(r, "N sigma:")
                theory_value = self.value_dict[product_symbol]
                theory_error = self.value_dict[self.error_dict[product_symbol]]
                output = n_sigma_format(theory_value, theory_error, res, er)
                self.insert_new_line(r, output)
    



#document = Document()

# p = document.add_paragraph()
# r = p.add_run()
# r.add_text('Good Morning every body,This is my ')
# r.add_picture('/tmp/foo.jpg')
# r.add_text(' do you like it?')

# document.save('demo.docx')

if __name__ == "__main__":
    da = ducument_analyser(_OUT_FOLDER, _KEYWORDS_FILTER)
    da.generate_document()